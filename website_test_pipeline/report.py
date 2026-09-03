"""Build human-verifiable Word reports from a test run.

The report exists so a person can confirm the pass/fail numbers instead of
trusting the generator + runner blindly: every test is shown with the
assertions it made and the screenshot evidence captured at each step.
"""
from __future__ import annotations

import ast
import io
import json
import os
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

IMAGE_EXTS = {".png", ".jpg", ".jpeg"}
ATTACH_EXTS = {".webm", ".zip"}

# Full-page screenshots are ~1280px wide and often 4-5k tall; embedded raw they
# bloat a full-run .docx to hundreds of MB. Downscale + JPEG-encode so the
# document stays shareable while text on the page is still readable.
_MAX_IMG_WIDTH = 1000
_MAX_IMG_HEIGHT = 3600
_JPEG_QUALITY = 55


def _embeddable(path: Path):
    """Return a JPEG BytesIO for `path`, or the original path if PIL can't read it."""
    try:
        from PIL import Image

        with Image.open(path) as im:
            im = im.convert("RGB")
            scale = min(_MAX_IMG_WIDTH / im.width, _MAX_IMG_HEIGHT / im.height, 1.0)
            if scale < 1.0:
                im = im.resize((round(im.width * scale), round(im.height * scale)), Image.LANCZOS)
            buffer = io.BytesIO()
            im.save(buffer, format="JPEG", quality=_JPEG_QUALITY, optimize=True)
            buffer.seek(0)
            return buffer
    except Exception:
        return str(path)


# --------------------------------------------------------------------------- model

@dataclass
class TestOutcome:
    nodeid: str
    title: str
    url: str
    status: str                       # passed | failed | skipped | error
    duration: float = 0.0
    error: str | None = None
    assertions: list[str] = field(default_factory=list)
    evidence: list[str] = field(default_factory=list)     # step screenshots
    attachments: list[str] = field(default_factory=list)  # video / trace

    @property
    def passed(self) -> bool:
        return self.status == "passed"


@dataclass
class UrlReport:
    url: str
    spec_path: str | None = None
    generated_status: str | None = None
    outcomes: list[TestOutcome] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    coverage_ceiling: str | None = None  # set when behavioural coverage is legitimately capped (captcha / embed / thin)

    @property
    def total(self) -> int:
        return len(self.outcomes)

    @property
    def passed(self) -> int:
        return sum(o.passed for o in self.outcomes)

    @property
    def failed(self) -> int:
        return sum(o.status in {"failed", "error"} for o in self.outcomes)


@dataclass
class RunReport:
    base_url: str = ""
    model: str = ""
    started_at: str = ""
    finished_at: str = ""
    url_reports: list[UrlReport] = field(default_factory=list)

    @property
    def total(self) -> int:
        return sum(u.total for u in self.url_reports)

    @property
    def passed(self) -> int:
        return sum(u.passed for u in self.url_reports)

    @property
    def failed(self) -> int:
        return sum(u.failed for u in self.url_reports)

    @property
    def warnings(self) -> list[str]:
        out: list[str] = []
        for u in self.url_reports:
            out.extend(f"[{u.url}] {w}" for w in u.warnings)
        return out


# ----------------------------------------------------------------------- assertions

def assertions_for(spec_path: Path, test_title: str) -> list[str]:
    """Pull the expect(...) / assert lines of one test function out of its source."""
    try:
        tree = ast.parse(spec_path.read_text(encoding="utf-8"))
    except (OSError, SyntaxError):
        return []
    for node in ast.walk(tree):
        if isinstance(node, ast.FunctionDef) and node.name == test_title:
            picked: list[str] = []
            for child in ast.walk(node):
                if isinstance(child, ast.Call) and _is_expect(child):
                    picked.append(_unparse(child))
                elif isinstance(child, ast.Assert):
                    picked.append("assert " + _unparse(child.test))
            return _dedupe(picked)
    return []


def _unparse(node: ast.AST) -> str:
    try:
        return ast.unparse(node)
    except Exception:
        return ""


def _is_expect(call: ast.Call) -> bool:
    target = call.func
    while isinstance(target, ast.Attribute):
        target = target.value
    return isinstance(target, ast.Call) and isinstance(target.func, ast.Name) and target.func.id == "expect"


def _dedupe(values: list[str]) -> list[str]:
    seen, out = set(), []
    for v in values:
        if v and v not in seen:
            seen.add(v)
            out.append(v)
    return out


# ---------------------------------------------------------------------------- load

def _slug(nodeid: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "-", nodeid).strip("-") or "test"


def _pw_slug(nodeid: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", nodeid.lower()).strip("-")


def load_run(artifacts_dir: Path, tests_dir: Path, model: str = "") -> RunReport:
    results = json.loads((artifacts_dir / "test_results.json").read_text(encoding="utf-8"))
    manifest = _maybe_json(artifacts_dir / "run.json")

    generated = {url: info.get("status") for url, info in (manifest.get("urls") or {}).items()}
    specs = {url: info.get("spec") for url, info in (manifest.get("urls") or {}).items()}

    by_url: dict[str, UrlReport] = {}
    for row in results.get("tests", []):
        url = row.get("url") or "(unknown url)"
        report = by_url.setdefault(url, UrlReport(url=url))
        report.spec_path = _resolve_spec(specs.get(url), row.get("nodeid", ""), tests_dir)
        report.generated_status = generated.get(url)

        spec = Path(report.spec_path) if report.spec_path else None
        title = row.get("nodeid", "").split("::")[-1].split("[")[0]
        outcome = TestOutcome(
            nodeid=row.get("nodeid", ""),
            title=row.get("title") or title,
            url=url,
            status=row.get("status", "error"),
            duration=float(row.get("duration") or 0.0),
            error=row.get("error"),
            assertions=assertions_for(spec, title) if spec and spec.exists() else [],
            evidence=_find_images(artifacts_dir / "evidence" / _slug(row.get("nodeid", ""))),
            attachments=_find_attachments(artifacts_dir / "pw", _pw_slug(row.get("nodeid", ""))),
        )
        report.outcomes.append(outcome)

    # generated specs that produced no test rows at all
    for url, status in generated.items():
        if status == "generated" and url not in by_url:
            by_url[url] = UrlReport(url=url, spec_path=specs.get(url), generated_status=status)

    run = RunReport(
        base_url=_common_prefix([u for u in by_url]),
        model=model or manifest.get("model", ""),
        started_at=manifest.get("started_at", ""),
        finished_at=manifest.get("finished_at", results.get("finished_at", "")),
        url_reports=[by_url[k] for k in sorted(by_url)],
    )
    for report in run.url_reports:
        report.coverage_ceiling = _behaviour_ceiling(_load_inventory(artifacts_dir, report.url))
        validate_url(report)
    return run


_ACTION_ROLES = {"button", "checkbox", "radio", "combobox", "tab", "switch", "slider", "menuitem", "menuitemcheckbox"}
_CAPTCHA_RE = re.compile(r"captcha|recaptcha|hcaptcha|are you human|prove you|robot", re.I)

def _inv_slug(url: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", url.lower()).strip("-")[:100]

def _load_inventory(artifacts_dir: Path, url: str) -> dict | None:
    path = artifacts_dir / f"{_inv_slug(url)}.inventory.json"
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return None

def _count_page_actions(data: dict) -> int:
    """How many genuinely interactive things the explorer saw - form controls,
    buttons, revealed panels, form fields."""
    n = 0
    for c in data.get("controls") or []:
        if c.get("region") == "chrome" or c.get("hidden"):
            continue
        tag, role = c.get("tag"), (c.get("role") or "")
        if (tag in {"button", "select", "textarea"}
                or (tag == "input" and (c.get("type") or "text") != "hidden")
                or role in _ACTION_ROLES):
            n += 1
    n += len(data.get("revealed") or [])
    for f in data.get("forms") or []:
        n += min(len(f.get("fields") or []), 3)
    return n

def _behaviour_ceiling(data: dict | None) -> str | None:
    """A reason the page's behavioural coverage is legitimately capped - so a file
    of mostly visibility checks is honest, not lazy. None means no excuse."""
    if data is None:
        return None
    haystack = json.dumps(data.get("controls") or [], ensure_ascii=False) + " " \
        + json.dumps(data.get("forms") or [], ensure_ascii=False) + " " \
        + (data.get("accessibility") or "")
    if _CAPTCHA_RE.search(haystack):
        return "the form is CAPTCHA-protected, so a real end-to-end submit cannot be automated"
    if data.get("embeds") and _count_page_actions(data) <= 3:
        return "the page is built around a third-party embed (map / media) with no driveable DOM"
    actions = _count_page_actions(data)
    if actions <= 2:
        return f"only ~{actions} interactive control(s) observed"
    return None

def _maybe_json(path: Path) -> dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}


def _guess_spec(nodeid: str, tests_dir: Path) -> str | None:
    stem = nodeid.split("::")[0]
    candidate = tests_dir / Path(stem).name
    return str(candidate) if candidate.exists() else None


def _resolve_spec(stored: str | None, nodeid: str, tests_dir: Path) -> str | None:
    """Prefer a stored path that still exists, else find the spec by name in tests_dir."""
    if stored and Path(stored).exists():
        return stored
    return _guess_spec(nodeid, tests_dir) or stored


def _find_images(directory: Path) -> list[str]:
    if not directory.is_dir():
        return []
    return [str(p) for p in sorted(directory.iterdir()) if p.suffix.lower() in IMAGE_EXTS]


def _find_attachments(pw_dir: Path, slug: str) -> list[str]:
    if not pw_dir.is_dir():
        return []
    out: list[str] = []
    for child in pw_dir.iterdir():
        if child.is_dir() and slug and slug in child.name:
            out.extend(str(p) for p in sorted(child.iterdir()) if p.suffix.lower() in ATTACH_EXTS)
    return out


def _common_prefix(urls: list[str]) -> str:
    if not urls:
        return ""
    first = urls[0]
    for i, ch in enumerate(first):
        if any(len(u) <= i or u[i] != ch for u in urls):
            return first[:i]
    return first


# ------------------------------------------------------------------------ validate

_BEHAVIOURAL_ASSERT = re.compile(
    r"to_have_value|to_have_url|to_contain_text|to_have_text|to_be_enabled|to_be_disabled|"
    r"to_be_checked|to_have_attribute|to_have_class|to_have_count\(\s*[2-9]"
)


def _title_calls_action(spec_path: Path | None, title: str) -> bool:
    """True if the test function performs a real action - it calls
    action_evidence(...) (which does a click/fill/select and then verifies a
    post-state). Such a test is behavioural regardless of which assertion the
    verify callback uses."""
    if not spec_path:
        return False
    try:
        tree = ast.parse(Path(spec_path).read_text(encoding="utf-8"))
    except (OSError, SyntaxError):
        return False
    for node in ast.walk(tree):
        if isinstance(node, ast.FunctionDef) and node.name == title:
            return any(
                isinstance(c, ast.Call) and isinstance(c.func, ast.Name) and c.func.id == "action_evidence"
                for c in ast.walk(node)
            )
    return False


def validate_url(report: UrlReport) -> None:
    warnings = report.warnings
    if report.generated_status == "generated" and report.total == 0:
        warnings.append("spec was generated but no tests executed")
    for outcome in report.outcomes:
        if outcome.passed and not outcome.evidence:
            warnings.append(f"{outcome.title}: PASSED with no screenshot evidence")
        if outcome.status in {"failed", "error"} and not outcome.attachments and not outcome.evidence:
            warnings.append(f"{outcome.title}: FAILED with no evidence or trace recorded")
        if outcome.passed and not outcome.assertions:
            warnings.append(f"{outcome.title}: PASSED with no detectable assertion (trivial test)")
    behavioural = sum(
        1 for o in report.outcomes
        if any(_BEHAVIOURAL_ASSERT.search(a) for a in o.assertions)
        or _title_calls_action(report.spec_path, o.title)
    )
    if report.total >= 4 and behavioural <= report.total // 4:
        if report.coverage_ceiling:
            # The page legitimately can't be driven further - a file of visibility
            # checks IS the honest coverage here, not a lazy one.
            warnings.append(
                f"limited coverage is expected here: {report.coverage_ceiling} "
                f"({behavioural}/{report.total} behavioural) - not a coverage gap"
            )
        else:
            warnings.append(
                f"{behavioural}/{report.total} tests exercise behaviour (click/fill/submit/navigate + verify); "
                "the rest only assert an element is visible - shallow coverage for this page"
            )


# --------------------------------------------------------------------------- render

def _heading_color(document, text: str, level: int, rgb: tuple[int, int, int] | None = None) -> None:
    heading = document.add_heading(text, level)
    if rgb:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*rgb)


def _status_line(document, outcome: TestOutcome) -> None:
    para = document.add_paragraph()
    run = para.add_run(f"Status: {outcome.status.upper()}")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x7F, 0x37) if outcome.passed else RGBColor(0xB3, 0x26, 0x1A)
    para.add_run(f"    Duration: {outcome.duration:.2f}s")


def _add_hyperlink(paragraph, target: str, text: str) -> None:
    r_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); props.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); props.append(underline)
    run.append(props)
    node = OxmlElement("w:t"); node.text = text; run.append(node)
    link.append(run)
    paragraph._p.append(link)


def _rel(target: str, base: Path | None) -> str:
    if base is None:
        return target
    try:
        return os.path.relpath(target, base).replace(os.sep, "/")
    except ValueError:
        return target


def _render_outcome(document, outcome: TestOutcome, *, embed: bool = True, link_base: Path | None = None) -> None:
    document.add_heading(outcome.title.replace("_", " "), 2)
    _status_line(document, outcome)

    if outcome.assertions:
        verified = outcome.status == "passed"
        document.add_heading(
            "Assertions verified" if verified else "Assertions in this test (test did not pass - see failure detail)", 3)
        for line in outcome.assertions:
            document.add_paragraph(line, style="List Bullet")

    document.add_heading("Browser evidence", 3)
    if not outcome.evidence:
        para = document.add_paragraph("No screenshot evidence was captured for this test.")
        para.runs[0].italic = True
    for image in outcome.evidence:
        path = Path(image)
        if not path.is_file():
            continue
        if embed:
            try:
                document.add_picture(_embeddable(path), width=Inches(6.0))
            except Exception as exc:  # unreadable / truncated screenshot
                document.add_paragraph(f"(could not embed {path.name}: {exc})")
                continue
            caption = document.add_paragraph(path.stem.replace("-", " "))
            caption.runs[0].italic = True
            caption.runs[0].font.size = Pt(9)
        else:
            para = document.add_paragraph(f"{path.stem.replace('-', ' ')}  —  ", style="List Bullet")
            _add_hyperlink(para, _rel(str(path), link_base), path.name)

    if outcome.status in {"failed", "error"}:
        document.add_heading("Failure detail", 3)
        block = document.add_paragraph(outcome.error or "(no traceback captured)")
        block.runs[0].font.name = "Consolas"
        block.runs[0].font.size = Pt(8)
        for attachment in outcome.attachments:
            para = document.add_paragraph("Attachment: ")
            para.runs[0].font.size = Pt(9)
            _add_hyperlink(para, _rel(attachment, link_base), Path(attachment).name)


def _summary_table(document, reports: list[UrlReport]) -> None:
    table = document.add_table(rows=1, cols=4)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for cell, label in zip(table.rows[0].cells, ("URL", "Tests", "Passed", "Failed")):
        cell.paragraphs[0].add_run(label).bold = True
    for report in reports:
        cells = table.add_row().cells
        cells[0].text = report.url
        cells[1].text = str(report.total)
        cells[2].text = str(report.passed)
        cells[3].text = str(report.failed)


def _metadata(document, run: RunReport, scope: str) -> None:
    document.add_paragraph(f"Scope: {scope}")
    document.add_paragraph(f"Model: {run.model or 'unknown'}")
    document.add_paragraph(f"Run started: {run.started_at or 'unknown'}")
    document.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).isoformat(timespec='seconds')}"
    )
    document.add_paragraph(
        f"Totals: {run.total} tests | {run.passed} passed | {run.failed} failed"
    )


def _warnings_section(document, warnings: list[str]) -> None:
    document.add_heading("Validation warnings", 1)
    if not warnings:
        document.add_paragraph("None. Every test has evidence and at least one assertion.")
        return
    document.add_paragraph(
        "These items mean the numbers above cannot be fully trusted from the "
        "document alone and need a manual look:"
    )
    for warning in warnings:
        para = document.add_paragraph(warning, style="List Bullet")
        para.runs[0].font.color.rgb = RGBColor(0xB3, 0x26, 0x1A)


def build_url_docx(run: RunReport, report: UrlReport, destination: Path) -> None:
    document = Document()
    document.add_heading("Website Test Evidence Report", 0)
    document.add_heading(report.url, 1)
    _metadata(document, run, scope=f"single URL ({report.url})")
    _summary_table(document, [report])
    _warnings_section(document, report.warnings)
    for outcome in report.outcomes:
        document.add_page_break()
        _render_outcome(document, outcome)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def build_combined_docx(run: RunReport, destination: Path) -> None:
    document = Document()
    document.add_heading("Website Test Evidence Report — Full Run", 0)
    _metadata(document, run, scope="all URLs")
    document.add_paragraph(
        "Screenshots are embedded beneath each test. Traces and videos remain "
        "linked; open this report from beside the artifacts/ folder so those links resolve."
    )
    _summary_table(document, run.url_reports)
    _warnings_section(document, run.warnings)
    link_base = destination.parent
    for report in run.url_reports:
        document.add_page_break()
        document.add_heading(report.url, 1)
        _summary_table(document, [report])
        if report.warnings:
            _warnings_section(document, report.warnings)
        for outcome in report.outcomes:
            document.add_page_break()
            _render_outcome(document, outcome, embed=True, link_base=link_base)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def name_for(url: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", url.lower()).strip("-")[:100] or "url"


def create_report(
    artifacts_dir: Path,
    tests_dir: Path,
    out_dir: Path,
    *,
    model: str = "",
    combined: bool = False,
) -> RunReport:
    run = load_run(artifacts_dir, tests_dir, model=model)
    out_dir.mkdir(parents=True, exist_ok=True)
    for report in run.url_reports:
        build_url_docx(run, report, out_dir / f"{name_for(report.url)}.docx")
    if combined:
        build_combined_docx(run, out_dir / "full-report.docx")
    return run
