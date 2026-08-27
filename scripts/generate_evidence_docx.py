import json, os, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
def clean_text(value):
    """Remove characters that WordprocessingML cannot store in XML text nodes."""
    if value is None: return ''
    import re
    text = re.sub(r'\x1b\[[0-?]*[ -/]*[@-~]', '', str(value))
    return ''.join(ch for ch in text if ch in '\t\n\r' or ord(ch) >= 32)
def cell_text(cell, text, bold=False, color=None):
    cell.text = ''; r = cell.paragraphs[0].add_run(clean_text(text)); r.bold = bold; r.font.size = Pt(9)
    if color: r.font.color.rgb = RGBColor(*color)
def main(src, dest):
    with open(src, encoding='utf-8') as fh: data = json.load(fh)
    if not isinstance(data, dict) or not isinstance(data.get('results', []), list): raise ValueError('Invalid evidence JSON schema: results must be a list')
    results = data.get('results', []); doc = Document(); sec = doc.sections[0]; sec.top_margin = sec.bottom_margin = Inches(.65); sec.left_margin = sec.right_margin = Inches(.7); doc.styles['Normal'].font.name = 'Aptos'; doc.styles['Normal'].font.size = Pt(10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run('Website Test Evidence Report'); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(31,78,121)
    p = doc.add_paragraph('Generated ' + clean_text(data.get('generatedAt',''))); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    passed = sum(x['status'] == 'passed' for x in results); skipped = sum(x['status'] == 'skipped' for x in results); failed = len(results) - passed - skipped; doc.add_heading('Executive summary', 1); doc.add_paragraph(f"{len(results)} tests executed: {passed} passed, {failed} failed, {skipped} not testable/skipped. Only named action evidence is included; generic end-of-test screenshots are excluded.")
    t = doc.add_table(rows=1, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c,h in zip(t.rows[0].cells,['Test','Status','Duration','Screenshots']): cell_text(c,h,True,(255,255,255))
    for x in results:
        if not isinstance(x, dict): continue
        cells=t.add_row().cells; cell_text(cells[0],x['title']); cell_text(cells[1],x['status'].upper(),True,(0,128,0) if x['status']=='passed' else (192,0,0)); cell_text(cells[2],f"{x.get('durationMs',0)} ms"); cell_text(cells[3],len(x.get('attachments',[])))
    for i,x in enumerate(results,1):
        doc.add_page_break(); doc.add_heading(clean_text(f"{i}. {x['title']}"),1); doc.add_paragraph(clean_text(f"Status: {x['status'].upper()} | Duration: {x.get('durationMs',0)} ms"))
        if x.get('error'): doc.add_heading('What happened',2); doc.add_paragraph(clean_text(x['error'].get('message') or 'The test failed without an error message.'))
        doc.add_heading('Browser evidence',2)
        if not x.get('attachments'): doc.add_paragraph('No verified action evidence was captured for this test. This is expected for observation-only or NOT TESTABLE tests.')
        for a in x.get('attachments',[]):
            if isinstance(a, dict) and a.get('path') and os.path.isfile(a['path']): doc.add_picture(a['path'],width=Inches(6.4)); cap=doc.add_paragraph(clean_text(a.get('name','Evidence'))); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
            elif isinstance(a, dict): doc.add_paragraph('Missing evidence attachment: ' + clean_text(a.get('name','unnamed')))
    os.makedirs(os.path.dirname(dest) or '.',exist_ok=True); doc.save(dest)
if __name__ == '__main__': main(sys.argv[1],sys.argv[2])
